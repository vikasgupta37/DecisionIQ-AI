import csv
import io
import json
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional
from fastapi import UploadFile
from backend.app.core.exceptions import DecisionIQException


# Allowed file extensions and their MIME type mappings
ALLOWED_EXTENSIONS = {
    ".csv": "text/csv",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pdf": "application/pdf",
    ".json": "application/json",
    ".txt": "text/plain",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# Maximum file size: 50 MB
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024


@dataclass
class FileMetadata:
    """Structured metadata extracted from an uploaded file."""
    filename: str
    file_type: str
    file_size: int
    row_count: Optional[int] = None
    column_names: Optional[List[str]] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    extra: Dict[str, Any] = field(default_factory=dict)


class FileProcessorService:
    """
    Validates uploaded files and extracts format-specific metadata.
    Supports CSV, Excel, PDF, JSON, TXT, and DOCX formats.
    """

    def validate_file(self, file: UploadFile) -> None:
        """
        Validates file type against allowed extensions and file size
        against the configured maximum.
        Raises DecisionIQException on validation failure.
        """
        filename = file.filename or ""
        ext = self._get_extension(filename)

        if ext not in ALLOWED_EXTENSIONS:
            raise DecisionIQException(
                message=f"File type '{ext}' is not supported. Allowed: {list(ALLOWED_EXTENSIONS.keys())}",
                status_code=400,
                code="INVALID_FILE_TYPE",
            )

        # Read content to check size, then reset pointer
        content = file.file.read()
        file_size = len(content)
        file.file.seek(0)

        if file_size > MAX_FILE_SIZE_BYTES:
            raise DecisionIQException(
                message=f"File size {file_size} bytes exceeds maximum allowed size of {MAX_FILE_SIZE_BYTES} bytes.",
                status_code=400,
                code="FILE_TOO_LARGE",
            )

        if file_size == 0:
            raise DecisionIQException(
                message="Uploaded file is empty.",
                status_code=400,
                code="EMPTY_FILE",
            )

    def extract_metadata(self, file: UploadFile) -> FileMetadata:
        """
        Extracts format-specific metadata from the uploaded file.
        Must be called after validate_file().
        """
        filename = file.filename or "unknown"
        ext = self._get_extension(filename)
        content = file.file.read()
        file_size = len(content)
        file.file.seek(0)

        metadata = FileMetadata(
            filename=filename,
            file_type=ext.lstrip("."),
            file_size=file_size,
        )

        try:
            if ext == ".csv":
                self._extract_csv_metadata(content, metadata)
            elif ext in (".xlsx", ".xls"):
                self._extract_excel_metadata(content, metadata)
            elif ext == ".json":
                self._extract_json_metadata(content, metadata)
            elif ext == ".pdf":
                self._extract_pdf_metadata(content, metadata)
            elif ext == ".txt":
                self._extract_txt_metadata(content, metadata)
            elif ext == ".docx":
                self._extract_docx_metadata(content, metadata)
        except Exception:
            # If metadata extraction fails, we still proceed with basic metadata
            pass

        return metadata

    def _get_extension(self, filename: str) -> str:
        """Extracts the lowercase file extension."""
        import os
        return os.path.splitext(filename)[1].lower()

    def _extract_csv_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts row count and column names from CSV content."""
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if rows:
            metadata.column_names = rows[0]
            metadata.row_count = max(len(rows) - 1, 0)  # Exclude header row

    def _extract_excel_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts row count and column names from Excel content."""
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        ws = wb.active
        if ws:
            rows = list(ws.iter_rows(values_only=True))
            if rows:
                metadata.column_names = [str(c) if c else "" for c in rows[0]]
                metadata.row_count = max(len(rows) - 1, 0)
            metadata.extra["sheet_count"] = len(wb.sheetnames)
            metadata.extra["sheet_names"] = wb.sheetnames
        wb.close()

    def _extract_json_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts record count and field names from JSON content."""
        text = content.decode("utf-8", errors="replace")
        data = json.loads(text)
        if isinstance(data, list):
            metadata.row_count = len(data)
            if data and isinstance(data[0], dict):
                metadata.column_names = list(data[0].keys())
        elif isinstance(data, dict):
            metadata.row_count = 1
            metadata.column_names = list(data.keys())

    def _extract_pdf_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts page count from PDF content."""
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        metadata.page_count = len(reader.pages)

    def _extract_txt_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts word count and line count from plain text content."""
        text = content.decode("utf-8", errors="replace")
        words = text.split()
        metadata.word_count = len(words)
        metadata.row_count = text.count("\n") + (1 if text and not text.endswith("\n") else 0)

    def _extract_docx_metadata(self, content: bytes, metadata: FileMetadata) -> None:
        """Extracts word count and paragraph count from DOCX content."""
        from docx import Document
        doc = Document(io.BytesIO(content))
        full_text = " ".join(p.text for p in doc.paragraphs)
        metadata.word_count = len(full_text.split())
        metadata.page_count = len(doc.paragraphs)


file_processor_service = FileProcessorService()
